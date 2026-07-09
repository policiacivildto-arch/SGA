from io import BytesIO, StringIO
from pathlib import Path
import re
import unicodedata
from datetime import date, datetime

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from django.conf import settings
from django.core.exceptions import ValidationError
from django.core.mail import EmailMessage
from django.core.management import call_command
from django.core.validators import validate_email
from django.db.models import Q, Sum
from django.http import HttpResponse
from rest_framework import status, viewsets
from rest_framework.authentication import BasicAuthentication, SessionAuthentication
from rest_framework.decorators import action
from rest_framework.permissions import IsAdminUser
from rest_framework.response import Response
from rest_framework.views import APIView
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from openpyxl import Workbook

from .models import (
    Arma,
    Cautela,
    Compra,
    Delegacia,
    Departamento,
    Fornecedor,
    Item,
    Lotacao,
    Movimento,
    OpcaoMenu,
    Patrimonio,
    Policial,
    Servico,
    UsuarioSistema,
    ITEM_STATUS_DISPONIVEL,
)
from .serializers import (
    ArmaSerializer,
    CautelaSerializer,
    CompraSerializer,
    DelegaciaSerializer,
    DepartamentoSerializer,
    FornecedorSerializer,
    ItemSerializer,
    LotacaoSerializer,
    MovimentoSerializer,
    OpcaoMenuSerializer,
    PatrimonioSerializer,
    PolicialSerializer,
    ServicoSerializer,
    UsuarioSistemaSerializer,
)


class BaseViewSet(viewsets.ModelViewSet):
    filter_map = {}

    def get_queryset(self):
        qs = super().get_queryset()

        for query_param, model_field in self.filter_map.items():
            value = self.request.query_params.get(query_param)

            if value is None or value == "":
                continue

            # Converte strings para boolean
            if isinstance(value, str):
                if value.lower() == "true":
                    value = True
                elif value.lower() == "false":
                    value = False

            qs = qs.filter(**{model_field: value})

        return qs


def normalize_text(value):
    text = (value or "").strip().lower()
    if not text:
        return ""
    text = unicodedata.normalize("NFKD", text)
    return "".join(ch for ch in text if not unicodedata.combining(ch))


DEFAULT_PLACEHOLDER = "Não informado"


class LotacaoViewSet(BaseViewSet):
    queryset = Lotacao.objects.all()
    serializer_class = LotacaoSerializer
    search_fields = ["depto", "nome", "cidade", "resp"]
    ordering_fields = ["depto", "nome", "cidade", "criado_em"]
    filter_map = {"depto": "depto"}


class DepartamentoViewSet(BaseViewSet):
    queryset = Departamento.objects.all()
    serializer_class = DepartamentoSerializer
    search_fields = ["nome", "sigla"]
    ordering_fields = ["nome", "sigla", "criado_em"]
    filter_map = {"ativo": "ativo"}


class DelegaciaViewSet(BaseViewSet):
    queryset = Delegacia.objects.select_related("departamento").all()
    serializer_class = DelegaciaSerializer
    search_fields = ["nome", "codigo", "cidade", "departamento__nome"]
    ordering_fields = ["nome", "codigo", "cidade", "criado_em"]
    filter_map = {"ativo": "ativo", "departamento": "departamento_id"}


class UsuarioSistemaViewSet(BaseViewSet):
    queryset = UsuarioSistema.objects.select_related("policial", "departamento", "delegacia").all()
    serializer_class = UsuarioSistemaSerializer
    search_fields = ["username", "nome", "cargo", "policial__nome", "departamento__nome", "delegacia__nome"]
    ordering_fields = ["username", "nome", "cargo", "criado_em"]
    filter_map = {"ativo": "ativo", "departamento": "departamento_id", "delegacia": "delegacia_id"}


class PolicialViewSet(BaseViewSet):
    queryset = Policial.objects.all()
    serializer_class = PolicialSerializer
    search_fields = ["matricula", "nome", "cargo", "depto", "lotacao"]
    ordering_fields = ["matricula", "nome", "depto", "lotacao", "criado_em"]
    filter_map = {"depto": "depto", "lotacao": "lotacao"}

    @staticmethod
    def _sync_lotacao(policial):
        depto = (policial.depto or "").strip()
        lotacao_nome = (policial.lotacao or "").strip()
        if not depto or not lotacao_nome:
            return

        Lotacao.objects.get_or_create(
            depto=depto,
            nome=lotacao_nome,
            defaults={
                "cidade": "N/D",
                "resp": "",
                "tel": "",
                "end": "",
            },
        )

    def perform_create(self, serializer):
        policial = serializer.save()
        self._sync_lotacao(policial)

    def perform_update(self, serializer):
        super().perform_update(serializer)
        self._sync_lotacao(serializer.instance)


class FornecedorViewSet(BaseViewSet):
    queryset = Fornecedor.objects.all()
    serializer_class = FornecedorSerializer
    search_fields = ["nome", "cnpj", "contato", "categoria"]
    ordering_fields = ["nome", "categoria", "criado_em"]


class CompraViewSet(BaseViewSet):
    queryset = Compra.objects.select_related("fornecedor").all()
    serializer_class = CompraSerializer
    search_fields = ["descricao", "categoria", "marca", "modelo", "serie", "status", "fornecedor__nome"]
    ordering_fields = ["dt_aq", "valor_compra", "categoria", "criado_em"]
    filter_map = {"categoria": "categoria", "status": "status", "fornecedor": "fornecedor_id"}

    @staticmethod
    def _normalize_quantidades(compra):
        qtd_total = max(int(compra.qtd_total or 1), 1)
        qtd_disp = int(compra.qtd_disp if compra.qtd_disp is not None else qtd_total)
        qtd_disp = max(0, min(qtd_disp, qtd_total))
        qtd_min = max(int(compra.qtd_min or 1), 0)
        return qtd_total, qtd_disp, qtd_min

    @staticmethod
    def _is_categoria_armas(categoria):
        return (str(categoria or "").strip().lower() == "armas")

    @staticmethod
    def _normalize_serie_text(value):
        return str(value or "").replace(" ", "").upper()

    @classmethod
    def _create_arma_detalhe(cls, compra, item, serie):
        if not cls._is_categoria_armas(compra.categoria):
            return None

        serie_norm = cls._normalize_serie_text(serie)

        return Arma.objects.create(
            item=item,
            tipo=compra.tipo,
            marca=compra.marca,
            modelo=compra.modelo,
            calibre=compra.calibre,
            comprimento_cano=compra.comprimento_cano,
            quantidade_carregadores=int(compra.quantidade_carregadores or 0),
            capacidade=int(compra.capacidade or 0),
            numero_serie=serie_norm,
        )

    @classmethod
    def _create_item_lote(cls, compra, qtd_total, qtd_disp, qtd_min):
        serie_norm = cls._normalize_serie_text(compra.serie)
        item = Item.objects.create(
            categoria=compra.categoria,
            marca=compra.marca,
            descricao=compra.descricao,
            serie=serie_norm,
            qtd_total=qtd_total,
            qtd_disp=qtd_disp,
            qtd_min=qtd_min,
            status=ITEM_STATUS_DISPONIVEL,
            fornecedor=compra.fornecedor,
            dt_aq=compra.dt_aq,
            valor_compra=compra.valor_compra,
            dt_val=compra.dt_val,
            obs=compra.obs,
        )
        cls._create_arma_detalhe(compra, item, serie_norm)
        return item

    @staticmethod
    def _validate_unidades(unidades, qtd_total):
        if not isinstance(unidades, list) or len(unidades) != qtd_total:
            return "Envie uma lista de unidades com a mesma quantidade da compra."

        for unidade in unidades:
            if not isinstance(unidade, dict):
                return "Formato invalido para unidades."

            serie = str(unidade.get("serie") or "").strip()
            patrimonio = str(unidade.get("patrimonio") or "").strip()
            if not serie and not patrimonio:
                return "Cada unidade deve ter ao menos serie ou numero de tombo."

        return ""

    @classmethod
    def _create_items_unidades(cls, compra, unidades):
        created_items = []
        for unidade in unidades:
            serie = cls._normalize_serie_text(unidade.get("serie"))
            patrimonio = str(unidade.get("patrimonio") or "").strip()
            item = Item.objects.create(
                categoria=compra.categoria,
                patrimonio=patrimonio,
                marca=compra.marca,
                descricao=compra.descricao,
                serie=serie,
                qtd_total=1,
                qtd_disp=1,
                qtd_min=0,
                status=ITEM_STATUS_DISPONIVEL,
                fornecedor=compra.fornecedor,
                dt_aq=compra.dt_aq,
                valor_compra=compra.valor_compra,
                dt_val=compra.dt_val,
                obs=compra.obs,
            )
            cls._create_arma_detalhe(compra, item, serie)
            created_items.append(item)
        return created_items

    @action(detail=True, methods=["post"], url_path="lancar-inventario")
    def lancar_inventario(self, request, pk=None):
        compra = self.get_object()
        status_norm = (compra.status or "").strip().lower()
        if "lanc" in status_norm:
            return Response({"detail": "Esta compra ja foi lancada no inventario."}, status=status.HTTP_400_BAD_REQUEST)

        categoria_norm = (compra.categoria or "").strip().lower()
        is_uniforme = categoria_norm == "uniformes"
        usar_unidades = bool(request.data.get("usar_unidades", False)) and not is_uniforme
        unidades = request.data.get("unidades") or []

        qtd_total, qtd_disp, qtd_min = self._normalize_quantidades(compra)

        if usar_unidades:
            validation_error = self._validate_unidades(unidades, qtd_total)
            if validation_error:
                return Response({"detail": validation_error}, status=status.HTTP_400_BAD_REQUEST)
            created_items = self._create_items_unidades(compra, unidades)
        else:
            created_items = [self._create_item_lote(compra, qtd_total, qtd_disp, qtd_min)]

        compra.status = "Lancada no Inventario"
        compra.save(update_fields=["status", "atualizado_em"])

        return Response(
            {
                "status": "ok",
                "compra_id": compra.id,
                "item_id": created_items[0].id,
                "item_ids": [item.id for item in created_items],
                "itens_criados": len(created_items),
                "detail": "Compra lancada no inventario com sucesso.",
            },
            status=status.HTTP_200_OK,
        )


class ItemViewSet(BaseViewSet):
    queryset = Item.objects.select_related("fornecedor").all()
    serializer_class = ItemSerializer
    search_fields = ["patrimonio", "descricao", "categoria", "marca", "serie", "status"]
    ordering_fields = ["descricao", "categoria", "qtd_total", "qtd_disp", "criado_em"]
    filter_map = {"categoria": "categoria", "status": "status"}

    @staticmethod
    def _inventario_headers():
        return [
            "ID",
            "Descricao",
            "Categoria",
            "Marca",
            "Serie",
            "Patrimonio",
            "Quantidade Total",
            "Quantidade Disponivel",
            "Quantidade Minima",
            "Status",
            "Fornecedor",
            "Data Aquisicao",
            "Data Validade",
            "Observacoes",
        ]

    @staticmethod
    def _inventario_row(item):
        return [
            item.id,
            item.descricao or "",
            item.categoria or "",
            item.marca or "",
            item.serie or "",
            item.patrimonio or "",
            item.qtd_total or 0,
            item.qtd_disp or 0,
            item.qtd_min or 0,
            item.status or "",
            item.fornecedor.nome if item.fornecedor else "",
            item.dt_aq.isoformat() if item.dt_aq else "",
            item.dt_val.isoformat() if item.dt_val else "",
            item.obs or "",
        ]

    @staticmethod
    def _apply_inventario_filters(queryset, categoria, status_filtro, search):
        qs = queryset
        if categoria:
            # Se a categoria for 'Armas', busca por 'Arma*' para incluir variações.
            if categoria.lower() == 'armas':
                qs = qs.filter(categoria__istartswith='arma')
            else:
                qs = qs.filter(categoria__icontains=categoria)
        if status_filtro:
            qs = qs.filter(status__iexact=status_filtro)
        if search:
            qs = qs.filter(
                Q(descricao__icontains=search)
                | Q(categoria__icontains=search)
                | Q(marca__icontains=search)
                | Q(serie__icontains=search)
                | Q(patrimonio__icontains=search)
                | Q(status__icontains=search)
                | Q(fornecedor__nome__icontains=search)
            )
        return qs

    @action(detail=False, methods=["get"], url_path="relatorio-xlsx")
    def relatorio_xlsx(self, request):
        base_qs = Item.objects.select_related("fornecedor").all().order_by("categoria", "descricao", "serie")

        categoria = (request.query_params.get("categoria") or "").strip()
        status_filtro = (request.query_params.get("status") or "").strip()
        search = (request.query_params.get("search") or "").strip()
        qs = self._apply_inventario_filters(base_qs, categoria, status_filtro, search)

        wb = Workbook()
        ws = wb.active
        ws.title = "Inventario"
        ws.append(self._inventario_headers())

        has_rows = False
        for item in qs.iterator():
            ws.append(self._inventario_row(item))
            has_rows = True

        if not has_rows:
            ws.append(["Sem registros para os filtros informados."])
        else:
            ws.auto_filter.ref = ws.dimensions

        output = BytesIO()
        wb.save(output)
        output.seek(0)

        filename = f"relatorio_inventario_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        response = HttpResponse(
            output.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response


class PatrimonioViewSet(BaseViewSet):
    queryset = Patrimonio.objects.select_related("departamento", "delegacia").all()
    serializer_class = PatrimonioSerializer
    search_fields = ["codigo", "descricao", "categoria", "departamento__nome", "delegacia__nome"]
    ordering_fields = ["codigo", "descricao", "categoria", "criado_em"]
    filter_map = {"ativo": "ativo", "departamento": "departamento_id", "delegacia": "delegacia_id"}


class ArmaViewSet(BaseViewSet):
    queryset = Arma.objects.select_related("item", "patrimonio").all()
    serializer_class = ArmaSerializer
    search_fields = ["tipo", "marca", "modelo", "calibre", "numero_serie", "item__descricao", "patrimonio__codigo"]
    ordering_fields = ["tipo", "marca", "modelo", "calibre", "criado_em"]
    filter_map = {"patrimonio": "patrimonio_id", "item": "item_id"}


class ServicoViewSet(BaseViewSet):
    queryset = Servico.objects.all()
    serializer_class = ServicoSerializer
    search_fields = [
        "codigo",
        "tipo",
        "armeiro",
        "status",
        "matricula",
        "policial_nome",
        "depto",
        "lotacao",
        "modelo",
        "serie",
    ]
    ordering_fields = ["codigo", "data_rec", "status", "criado_em"]
    filter_map = {"status": "status", "tipo": "tipo", "armeiro": "armeiro"}

    @action(detail=False, methods=["get"], url_path="next-code")
    def next_code(self, request):
        # A lógica foi movida para um método de classe no modelo
        return Response({"codigo": Servico.get_next_code()})

    def perform_create(self, serializer):
        policial_id = serializer.validated_data.get("policial_id")
        policial = None
        if policial_id:
            try:
                policial = Policial.objects.get(pk=policial_id)
            except Policial.DoesNotExist:
                # Lida com o caso de um ID de policial inválido, se necessário
                pass

        extra_data = {}
        if policial:
            # Associa o policial e preenche os campos denormalizados
            extra_data["policial"] = policial
            extra_data["matricula"] = policial.matricula
            extra_data["policial_nome"] = policial.nome
            extra_data["depto"] = policial.depto
            extra_data["lotacao"] = policial.lotacao

        serializer.save(**extra_data)

class CautelaViewSet(BaseViewSet):
    queryset = Cautela.objects.select_related(
        "policial",
        "item",
        "item__fornecedor",
        "item__arma_detalhe",
        "item__arma_detalhe__patrimonio",
    ).all()
    serializer_class = CautelaSerializer
    search_fields = [
        "numero",
        "policial_nome",
        "matricula",
        "depto",
        "lotacao",
        "item__descricao", # Busca no campo relacionado
        "serie",
        "status",
    ]
    ordering_fields = ["numero", "data_saida", "status", "criado_em"]
    filter_map = {"status": "status", "categoria": "categoria", "lotacao": "lotacao"}

    @staticmethod
    def _destinatario_label(cautela):
        return "Pessoal" if str(cautela.matricula or "").strip() else "Unidade"

    @staticmethod
    def _fortaleza_data_text():
        now = datetime.now()
        weekdays = [
            "segunda-feira",
            "terca-feira",
            "quarta-feira",
            "quinta-feira",
            "sexta-feira",
            "sabado",
            "domingo",
        ]
        months = [
            "janeiro",
            "fevereiro",
            "marco",
            "abril",
            "maio",
            "junho",
            "julho",
            "agosto",
            "setembro",
            "outubro",
            "novembro",
            "dezembro",
        ]
        return f"Fortaleza, {now.day} ({weekdays[now.weekday()]}) de {months[now.month - 1]} de {now.year}."

    @classmethod
    def _termos_responsabilidade(cls):
        return [
            "O(A) servidor(a) identificado(a) neste termo declara ter recebido o(s) material(is) descrito(s) acima em perfeito estado de conservacao e funcionamento, comprometendo-se a utiliza-lo(s) exclusivamente em atividades do servico, cabendo ainda:",
            "1. Zelar pela sua integridade, economia e conservacao dos bens do Estado, especialmente daqueles que lhes sejam entregues para guarda ou utilizacao (Art. 100, §2º, da Lei nº 12.124, de 06 de julho de 1993);",
            "2. Nao exibir desnecessariamente arma, distintivo ou algemas (Art. 103, item A, §4º, da Lei nº 12.124, de 06 de julho de 1993)",
            "3. Comunicar imediatamente qualquer ocorrencia envolvendo armamento sob sua responsabilidade, inclusive registrando boletim de ocorrencia",
            "4. Devolver, no dia da exoneracao ou demissao, ou quando solicitado pelo setor responsavel, os objetos pertencentes ao Estado que lhes foram entregues para sua guarda ou utilizacao (Art. 163, §1º, da Lei nº 12.124, de 06 de julho de 1993) nas condicoes em que o(s) recebeu.",
            "5. Devolver a arma quando do afastamento para aposentadoria (Provimento Correcional nº 02/2012 - CGD).",
            "Declaro para os devidos fins que detenho conhecimento tecnico para o manuseio do referido armamento.",
            cls._fortaleza_data_text(),
        ]

    @staticmethod
    def _cautela_header_lines():
        return [
            "GOVERNO DO ESTADO DO CEARA",
            "SECRETARIA DA SEGURANCA PUBLICA E DEFESA SOCIAL",
            "DELEGACIA GERAL DA POLICIA CIVIL",
            "DEPARTAMENTO TECNICO OPERACIONAL",
            "UNIDADE DE APOIO LOGISTICO",
        ]

    @staticmethod
    def _cautela_intro_text():
        return "Recebi da POLICIA CIVIL DO ESTADO DO CEARA, o ITEM abaixo especificado:"

    @staticmethod
    def _cautela_logo_path():
        logo_path = Path(__file__).resolve().parents[2] / "POLICIA CIVIL CE sem fundo.png"
        if logo_path.exists():
            return logo_path
        return None

    @staticmethod
    def _word_rows(cautela):
        def safe(value):
            return "" if value is None else str(value)

        item = cautela.item
        arma = getattr(item, "arma_detalhe", None)
        patrimonio = getattr(arma, "patrimonio", None) if arma else None
        fornecedor = getattr(item, "fornecedor", None)
        item_nome = cautela.item_desc if cautela.item_desc else getattr(item, "descricao", "")
        destinatario = CautelaViewSet._destinatario_label(cautela)
        emitido_em = datetime.now().strftime("%d/%m/%Y %H:%M:%S")

        return [
            ("Emitido em", emitido_em),
            ("Numero", safe(cautela.numero)),
            ("Status", safe(cautela.status)),
            ("Data de Saida", safe(cautela.data_saida)),
            ("Previsao de Devolucao", safe(cautela.data_prev)),
            ("Data de Devolucao", safe(cautela.data_dev)),
            ("Destinatario", destinatario),
            ("Policial", safe(cautela.policial_nome)),
            ("Matricula", safe(cautela.matricula)),
            ("Departamento", safe(cautela.depto)),
            ("Lotacao", safe(cautela.lotacao)),
            ("Item", safe(item_nome)),
            ("Categoria", safe(cautela.categoria)),
            ("Serie/Patrimonio", safe(cautela.serie)),
            ("Quantidade", safe(cautela.qtd)),
            ("Patrimonio (Item)", safe(getattr(item, "patrimonio", ""))),
            ("Marca (Item)", safe(getattr(item, "marca", ""))),
            ("Serie (Item)", safe(getattr(item, "serie", ""))),
            ("Data Aquisicao", safe(getattr(item, "dt_aq", ""))),
            ("Data Validade", safe(getattr(item, "dt_val", ""))),
            ("Valor Compra", safe(getattr(item, "valor_compra", ""))),
            ("Fornecedor", safe(getattr(fornecedor, "nome", ""))),
            ("Tipo da Arma", safe(getattr(arma, "tipo", ""))),
            ("Marca da Arma", safe(getattr(arma, "marca", ""))),
            ("Modelo da Arma", safe(getattr(arma, "modelo", ""))),
            ("Calibre", safe(getattr(arma, "calibre", ""))),
            ("Comprimento do Cano", safe(getattr(arma, "comprimento_cano", ""))),
            ("Capacidade", safe(getattr(arma, "capacidade", ""))),
            ("Qtd. Carregadores", safe(getattr(arma, "quantidade_carregadores", ""))),
            ("Numero de Serie da Arma", safe(getattr(arma, "numero_serie", ""))),
            ("Tombo Patrimonial", safe(getattr(patrimonio, "codigo", ""))),
            ("Observacoes", safe(cautela.obs)),
            ("Condicao na Devolucao", safe(cautela.condicao_dev)),
            ("Obs. Devolucao", safe(cautela.obs_dev)),
        ]

    @staticmethod
    def _render_word_document(rows):
        document = Document()
        
        # 1. Configura margens para padrão oficial (Ofício)
        for section in document.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.0)

        # 2. Insere a logo no topo
        logo_path = CautelaViewSet._cautela_logo_path()
        if logo_path:
            try:
                logo_paragraph = document.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Zera o espaçamento abaixo da imagem para que o texto suba
                logo_paragraph.paragraph_format.space_after = Pt(0)
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(str(logo_path), width=Cm(2.2)) # Tamanho ajustado
            except Exception:
                pass

        # 3. Insere as linhas do cabeçalho agrupadas logo abaixo da imagem
        header_lines = CautelaViewSet._cautela_header_lines()
        for i, line in enumerate(header_lines):
            paragraph = document.add_paragraph(line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Remove o espaçamento entre as linhas, exceto na última
            if i < len(header_lines) - 1:
                paragraph.paragraph_format.space_after = Pt(0)
            else:
                paragraph.paragraph_format.space_after = Pt(18) # Dá um respiro antes do título da Cautela
                
            run = paragraph.runs[0]
            run.bold = True
            run.font.name = 'Arial'  # Fonte sóbria
            run.font.size = Pt(10)   # Tamanho menor, típico de cabeçalho institucional

        # 4. Continua com o documento padrão
        heading = document.add_heading("CAUTELA", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        intro_paragraph = document.add_paragraph(CautelaViewSet._cautela_intro_text())
        intro_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph("")

        table = document.add_table(rows=0, cols=2)
        for label, value in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = str(label)
            row_cells[1].text = str(value)

        document.add_paragraph("")
        document.add_heading("4. Termos e Responsabilidade", level=2)
        for texto in CautelaViewSet._termos_responsabilidade():
            document.add_paragraph(texto)
            
        document.add_paragraph("")
        document.add_paragraph("_______________________________")
        document.add_paragraph("Recebedor(a)")
        document.add_paragraph("")
        document.add_paragraph("_______________________________")
        document.add_paragraph("Entregador(a)")
        
        return document
    
    @staticmethod
    def _word_filename(cautela):
        safe_number = re.sub(r"[^A-Za-z0-9_-]", "_", cautela.numero or f"cautela_{cautela.id}")
        return f"cautela_{safe_number}.docx"

    @staticmethod
    def _pdf_filename(cautela):
        safe_number = re.sub(r"[^A-Za-z0-9_-]", "_", cautela.numero or f"cautela_{cautela.id}")
        return f"cautela_{safe_number}.pdf"

    @staticmethod
    def _pdf_safe(value):
        if value is None:
            return ""
        return str(value)

    @staticmethod
    def _email_subject(cautela):
        numero = cautela.numero or f"#{cautela.id}"
        return f"Cautela {numero}"

    @classmethod
    def _email_body(cls, cautela):
        numero = cautela.numero or f"#{cautela.id}"
        return (
            "Prezados,\n\n"
            f"Segue em anexo o documento da cautela {numero}.\n"
            "Em caso de duvidas, favor contatar o setor responsavel.\n\n"
            "Mensagem enviada automaticamente pelo sistema SALT."
        )

    @classmethod
    def _build_email_attachment(cls, cautela, formato):
        formato_norm = str(formato or "pdf").strip().lower()
        if formato_norm in {"word", "docx"}:
            rows = cls._word_rows(cautela)
            document = cls._render_word_document(rows)
            output = BytesIO()
            document.save(output)
            output.seek(0)
            return (
                cls._word_filename(cautela),
                output.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        pdf_output = cls._build_pdf_document(cautela)
        return (
            cls._pdf_filename(cautela),
            pdf_output.getvalue(),
            "application/pdf",
        )

    @classmethod
    def _build_pdf_document(cls, cautela):
        from reportlab.lib.utils import ImageReader
        from reportlab.lib.styles import ParagraphStyle

        item = cautela.item
        arma = getattr(item, "arma_detalhe", None)
        patrimonio = getattr(arma, "patrimonio", None) if arma else None
        fornecedor = getattr(item, "fornecedor", None)
        policial = cautela.policial
        destinatario = cls._destinatario_label(cautela)

        styles = getSampleStyleSheet()
        styles["Heading3"].alignment = 1
        styles["Title"].alignment = 1

        # Cria um estilo institucional elegante para o cabeçalho (sem itálico e em negrito)
        estilo_cabecalho = ParagraphStyle(
            'CabecalhoOficial',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=10,
            alignment=1,  # Centralizado
            spaceAfter=2,
        )

        output = BytesIO()
        doc = SimpleDocTemplate(
            output,
            pagesize=A4,
            leftMargin=1.5 * cm,
            rightMargin=1.5 * cm,
            topMargin=1.5 * cm,
            bottomMargin=1.5 * cm,
        )

        story = []
        logo_path = cls._cautela_logo_path()
        if logo_path:
            try:
                # ImageReader pega as dimensões reais da imagem para travar a proporção
                img_reader = ImageReader(str(logo_path))
                iw, ih = img_reader.getSize()
                
                # Define a largura fixa (ex: 2.2 cm) e calcula a altura proporcional
                largura_logo = 2.2 * cm
                proporcao = ih / float(iw)
                altura_logo = largura_logo * proporcao
                
                logo_img = RLImage(str(logo_path), width=largura_logo, height=altura_logo)
                logo_img.hAlign = "CENTER"
                story.append(logo_img)
                story.append(Spacer(1, 0.2 * cm))
            except Exception:
                pass

        # Cabeçalho usando o novo estilo (mantendo os textos originais para não gerar erro de encoding)
        story.extend([
            Paragraph("GOVERNO DO ESTADO DO CEARA", estilo_cabecalho),
            Paragraph("SECRETARIA DA SEGURANCA PUBLICA E DEFESA SOCIAL", estilo_cabecalho),
            Paragraph("DELEGACIA GERAL DA POLICIA CIVIL", estilo_cabecalho),
            Paragraph("DEPARTAMENTO TECNICO OPERACIONAL", estilo_cabecalho),
            Paragraph("UNIDADE DE APOIO LOGISTICO", estilo_cabecalho),
            Spacer(1, 0.5 * cm),
            Paragraph("CAUTELA", styles["Title"]),
            Paragraph(cls._cautela_intro_text(), styles["Normal"]),
            Spacer(1, 0.4 * cm),
            Paragraph(f"Numero: {cls._pdf_safe(cautela.numero)}", styles["Normal"]),
            Paragraph(f"Emitido em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", styles["Normal"]),
            Spacer(1, 0.35 * cm),
            Paragraph("Dados da Cautela", styles["Heading3"]),
        ])

        dados_cautela = [
            ["Campo", "Valor"],
            ["Status", cls._pdf_safe(cautela.status)],
            ["Data de Saida", cls._pdf_safe(cautela.data_saida)],
            ["Previsao de Devolucao", cls._pdf_safe(cautela.data_prev) or "Sem prazo definido"],
            ["Data de Devolucao", cls._pdf_safe(cautela.data_dev)],
            ["Destinatario", destinatario],
            ["Departamento", cls._pdf_safe(cautela.depto)],
            ["Lotacao", cls._pdf_safe(cautela.lotacao)],
            ["Observacoes", cls._pdf_safe(cautela.obs)],
        ]

        tabela_cautela = Table(dados_cautela, colWidths=[5.0 * cm, 12.0 * cm])
        tabela_cautela.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.7, colors.black),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.extend([tabela_cautela, Spacer(1, 0.35 * cm), Paragraph("Dados do Policial", styles["Heading3"])])

        dados_policial = [
            ["Campo", "Valor"],
            ["Nome", cls._pdf_safe(cautela.policial_nome)],
            ["Cargo", cls._pdf_safe(getattr(policial, "cargo", ""))],
            ["E-mail", cls._pdf_safe(getattr(policial, "email", ""))],
            ["Lotacao", cls._pdf_safe(cautela.lotacao)],
            ["Matricula", cls._pdf_safe(cautela.matricula)],
        ]
        tabela_policial = Table(dados_policial, colWidths=[5.0 * cm, 12.0 * cm])
        tabela_policial.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.7, colors.black),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.extend([tabela_policial, Spacer(1, 0.35 * cm), Paragraph("Dados do Material", styles["Heading3"])])

        dados_material = [
            ["Campo", "Valor"],
            ["Item", cls._pdf_safe(cautela.item_desc or item.descricao)],
            ["Categoria", cls._pdf_safe(cautela.categoria)],
            ["Quantidade", cls._pdf_safe(cautela.qtd)],
            ["Serie/Patrimonio", cls._pdf_safe(cautela.serie)],
            ["Fornecedor", cls._pdf_safe(getattr(fornecedor, "nome", ""))],
            ["Tipo da Arma", cls._pdf_safe(getattr(arma, "tipo", ""))],
            ["Marca", cls._pdf_safe(getattr(arma, "marca", "") or getattr(item, "marca", ""))],
            ["Modelo", cls._pdf_safe(getattr(arma, "modelo", ""))],
            ["Calibre", cls._pdf_safe(getattr(arma, "calibre", ""))],
            ["Comprimento do Cano", cls._pdf_safe(getattr(arma, "comprimento_cano", ""))],
            ["Capacidade", cls._pdf_safe(getattr(arma, "capacidade", ""))],
            ["Qtd. Carregadores", cls._pdf_safe(getattr(arma, "quantidade_carregadores", ""))],
            ["Numero de Serie da Arma", cls._pdf_safe(getattr(arma, "numero_serie", ""))],
            ["Tombo Patrimonial", cls._pdf_safe(getattr(patrimonio, "codigo", ""))],
        ]
        tabela_material = Table(dados_material, colWidths=[5.0 * cm, 12.0 * cm])
        tabela_material.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.7, colors.black),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.extend([tabela_material, Spacer(1, 0.4 * cm)])

        story.append(Paragraph("4. Termos e Responsabilidade", styles["Heading3"]))
        for texto in cls._termos_responsabilidade():
            story.append(Paragraph(texto, styles["Normal"]))
            story.append(Spacer(1, 0.15 * cm))

        # Espaço vertical para a assinatura manual
        story.append(Spacer(1, 1.5 * cm)) 

        # Tabela invisível para deixar assinaturas lado a lado
        dados_assinaturas = [
            ["_______________________________", "_______________________________"],
            ["Recebedor(a)", "Entregador(a)"]
        ]
        
        tabela_assinaturas = Table(dados_assinaturas, colWidths=[9.0 * cm, 9.0 * cm])
        tabela_assinaturas.setStyle(TableStyle([
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
        ]))
        
        story.append(tabela_assinaturas)

        # Processamento final OBRIGATÓRIO (aqui que o arquivo corrompido é evitado)
        doc.build(story)
        output.seek(0)
        return output

    @action(detail=False, methods=["get"], url_path="next-number")
    def next_number(self, request):
        year = str(datetime.now().year)
        last_number = (
            Cautela.objects.filter(numero__startswith=f"CAU-{year}-")
            .order_by("-numero")
            .values_list("numero", flat=True)
            .first()
        )
        next_num = 1
        if last_number:
            try:
                next_num = int(last_number.split("-")[-1]) + 1
            except (ValueError, IndexError):
                next_num = 1
        return Response({"numero": f"CAU-{year}-{str(next_num).zfill(4)}"})

    @action(detail=True, methods=["post"], url_path="devolver")
    def devolver(self, request, pk=None):
        cautela = self.get_object()
        data_dev = request.data.get("data_dev")
        condicao_dev = request.data.get("condicao_dev", "")
        obs_dev = request.data.get("obs_dev", "")

        if not data_dev:
            return Response({"detail": "data_dev e obrigatorio."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # A lógica de devolução foi movida para um método de instância no modelo
            cautela.devolver_cautela(
                data_dev=data_dev,
                condicao_dev=condicao_dev,
                obs_dev=obs_dev,
            )
        except ValueError as exc:
            return Response({"detail": str(exc)}, status=status.HTTP_400_BAD_REQUEST)

        return Response(self.get_serializer(cautela).data)

    @action(detail=True, methods=["get"], url_path="documento-word")
    def documento_word(self, request, pk=None):
        cautela = self.get_object()
        rows = self._word_rows(cautela)
        document = self._render_word_document(rows)
        filename = self._word_filename(cautela)

        output = BytesIO()
        document.save(output)
        output.seek(0)

        response = HttpResponse(
            output.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    @action(detail=True, methods=["get"], url_path="documento-pdf")
    def documento_pdf(self, request, pk=None):
        cautela = self.get_object()
        filename = self._pdf_filename(cautela)
        output = self._build_pdf_document(cautela)

        response = HttpResponse(output.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    @action(detail=True, methods=["post"], url_path="enviar-email")
    def enviar_email(self, request, pk=None):
        cautela = self.get_object()

        formato = (request.data.get("formato") or "pdf").strip().lower()
        if formato not in {"pdf", "word", "docx"}:
            return Response(
                {"detail": "Formato invalido. Use 'pdf' ou 'word'."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        email_destino = (request.data.get("email") or "").strip()
        if not email_destino:
            email_destino = str(getattr(cautela.policial, "email", "") or "").strip()

        if not email_destino:
            return Response(
                {"detail": "Informe um e-mail destino ou cadastre e-mail no policial da cautela."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            validate_email(email_destino)
        except ValidationError:
            return Response({"detail": "E-mail destino invalido."}, status=status.HTTP_400_BAD_REQUEST)

        filename, attachment_bytes, content_type = self._build_email_attachment(cautela, formato)

        from_email = getattr(settings, "DEFAULT_FROM_EMAIL", "no-reply@salt.local")
        mensagem = EmailMessage(
            subject=self._email_subject(cautela),
            body=self._email_body(cautela),
            from_email=from_email,
            to=[email_destino],
        )
        mensagem.attach(filename, attachment_bytes, content_type)

        try:
            enviados = mensagem.send(fail_silently=False)
        except Exception as exc:
            return Response(
                {"detail": f"Falha ao enviar e-mail: {exc}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        if enviados < 1:
            return Response(
                {"detail": "Nenhum e-mail foi enviado pelo backend."},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        return Response(
            {
                "status": "ok",
                "detail": "Cautela enviada por e-mail com sucesso.",
                "email": email_destino,
                "formato": "word" if formato in {"word", "docx"} else "pdf",
                "arquivo": filename,
            },
            status=status.HTTP_200_OK,
        )

    @staticmethod
    def _xlsx_headers():
        return [
            "Numero",
            "Data Saida",
            "Data Prevista",
            "Data Devolucao",
            "Status",
            "Destinatario",
            "Policial",
            "Matricula",
            "Departamento",
            "Lotacao",
            "Item",
            "Categoria",
            "Tipo",
            "Numero de Serie",
            "Quantidade",
            "Observacoes",
        ]

    def _xlsx_row(self, cautela):
        destinatario = self._destinatario_label(cautela)
        tipo_arma = ""
        if (cautela.categoria or "").lower() == "armas":
            arma_detalhe = getattr(cautela.item, "arma_detalhe", None)
            if arma_detalhe:
                tipo_arma = arma_detalhe.tipo or ""
        return [
            cautela.numero or "",
            cautela.data_saida.isoformat() if cautela.data_saida else "",
            cautela.data_prev.isoformat() if cautela.data_prev else "",
            cautela.data_dev.isoformat() if cautela.data_dev else "",
            cautela.status or "",
            destinatario,
            cautela.policial_nome or "",
            cautela.matricula or "",
            cautela.depto or "",
            cautela.lotacao or "",
            cautela.item_desc or (cautela.item.descricao if cautela.item else ""),
            cautela.categoria or "",
            tipo_arma,
            cautela.serie or "",
            cautela.qtd or 0,
            cautela.obs or "",
        ]

    @staticmethod
    def _apply_xlsx_filters(queryset, categoria, status_filtro, search):
        qs = queryset
        if categoria:
            categoria_norm = str(categoria).strip().lower()
            if categoria_norm.startswith("arma"):
                qs = qs.filter(categoria__icontains="arma")
            else:
                qs = qs.filter(Q(categoria__iexact=categoria) | Q(categoria__icontains=categoria))
        if status_filtro:
            qs = qs.filter(status__iexact=status_filtro)
        if search:
            qs = qs.filter(
                Q(numero__icontains=search)
                | Q(policial_nome__icontains=search)
                | Q(matricula__icontains=search)
                | Q(depto__icontains=search)
                | Q(lotacao__icontains=search)
                | Q(item_desc__icontains=search)
                | Q(categoria__icontains=search)
                | Q(serie__icontains=search)
            )
        return qs

    @action(detail=False, methods=["get"], url_path="relatorio-xlsx")
    def relatorio_xlsx(self, request):
        base_qs = (
            Cautela.objects.select_related("item", "item__arma_detalhe")
            .all()
            .order_by("-data_saida", "-criado_em")
        )

        categoria = (request.query_params.get("categoria") or "").strip()
        status_filtro = (request.query_params.get("status") or "").strip()
        search = (request.query_params.get("search") or "").strip()
        qs = self._apply_xlsx_filters(base_qs, categoria, status_filtro, search)

        wb = Workbook()
        ws = wb.active
        ws.title = "Cautelas"
        ws.append(self._xlsx_headers())

        has_rows = False
        for c in qs.iterator():
            ws.append(self._xlsx_row(c))
            has_rows = True

        if not has_rows:
            ws.append([
                "Sem registros para os filtros informados.",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ])
        else:
            ws.auto_filter.ref = ws.dimensions

        output = BytesIO()
        wb.save(output)
        output.seek(0)

        filename = f"relatorio_cautelas_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        response = HttpResponse(
            output.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response


class MovimentoViewSet(BaseViewSet):
    queryset = Movimento.objects.select_related("item").all()
    serializer_class = MovimentoSerializer
    search_fields = [
        "tipo",
        "item_desc",
        "categoria",
        "serie",
        "policial",
        "matricula",
        "lotacao",
        "num_cautela",
        "obs",
    ]
    ordering_fields = ["data", "tipo", "criado_em"]
    filter_map = {"tipo": "tipo", "categoria": "categoria", "lotacao": "lotacao"}


class OpcaoMenuViewSet(BaseViewSet):
    queryset = OpcaoMenu.objects.all()
    serializer_class = OpcaoMenuSerializer
    search_fields = ["grupo", "valor", "rotulo"]
    ordering_fields = ["grupo", "ordem", "rotulo", "criado_em"]
    filter_map = {"ativo": "ativo"}

    def get_queryset(self):
        qs = super().get_queryset()
        if self.request.query_params.get("grupo"):
            grupos = self.request.query_params.get("grupo").split(",")
            qs = qs.filter(grupo__in=grupos)
        return qs


class DashboardEstoqueResumoView(APIView):
    @staticmethod
    def _build_locais():
        return [
            {
                "departamento": lot["depto"] or "",
                "delegacia": lot["nome"] or "",
                "seccional": lot["cidade"] or "",
            }
            for lot in Lotacao.objects.values("depto", "nome", "cidade").order_by("depto", "nome")
        ]

    @staticmethod
    def _build_categorias():
        return list(Item.objects.values_list("categoria", flat=True).distinct().order_by("categoria"))

    @staticmethod
    def _build_totais(itens_base):
        totais = itens_base.aggregate(total=Sum("qtd_total"), disp=Sum("qtd_disp"))
        return {
            "total": int(totais["total"] or 0),
            "disp": int(totais["disp"] or 0),
            "registros_itens": itens_base.count(),
        }

    @staticmethod
    def _build_ranked_totals(itens_base, field_name):
        return [
            [row[field_name] or DEFAULT_PLACEHOLDER, int(row["total"] or 0)]
            for row in itens_base.values(field_name).annotate(total=Sum("qtd_total")).order_by("-total", field_name)[:8]
        ]

    @staticmethod
    def _build_by_modelo(itens_base):
        # Filtra apenas itens da categoria 'Armas' e busca os modelos do detalhe
        armas_base = itens_base.filter(categoria__icontains="arma").select_related("arma_detalhe")
        return [
            [row["arma_detalhe__modelo"] or DEFAULT_PLACEHOLDER, int(row["total"] or 0)]
            for row in armas_base.values("arma_detalhe__modelo").annotate(total=Sum("qtd_total")).order_by("-total", "arma_detalhe__modelo")[:10]
            if row["arma_detalhe__modelo"]
        ]

    @staticmethod
    def _build_validade(itens_base):
        hoje = date.today()
        qtd_vencidos = 0
        qtd_vence30 = 0
        qtd_vence60 = 0
        qtd_sem_validade = 0
        itens_criticos = []
        por_ano = {}

        for item in itens_base.values("id", "descricao", "categoria", "dt_val", "qtd_total", "qtd_disp"):
            qtd_base = int(item["qtd_disp"] if item["qtd_disp"] is not None else item["qtd_total"] or 0)
            dt_val = item["dt_val"]

            if not dt_val:
                qtd_sem_validade += qtd_base
                continue

            ano = dt_val.year
            por_ano[ano] = por_ano.get(ano, 0) + qtd_base

            dias_para_vencer = (dt_val - hoje).days
            situacao = "Dentro da validade"

            if dias_para_vencer < 0:
                situacao = "Vencido"
                qtd_vencidos += qtd_base
            elif dias_para_vencer <= 365:
                situacao = "Vence em ate 1 ano"
                qtd_vence30 += qtd_base
            elif dias_para_vencer <= 730:
                situacao = "Vence em ate 2 anos"
                qtd_vence60 += qtd_base

            if dias_para_vencer <= 730:
                itens_criticos.append(
                    {
                        "id": item["id"],
                        "descricao": item["descricao"] or "—",
                        "categoria": item["categoria"] or DEFAULT_PLACEHOLDER,
                        "dt_val": dt_val,
                        "qtd_disp": qtd_base,
                        "diasParaVencer": dias_para_vencer,
                        "anosParaVencer": round(dias_para_vencer / 365, 2),
                        "situacao": situacao,
                    }
                )

        itens_criticos.sort(key=lambda item: (item["diasParaVencer"], item["descricao"].lower()))
        validade_por_ano = [
            {"ano": ano, "quantidade": int(qtd)}
            for ano, qtd in sorted(por_ano.items(), key=lambda pair: pair[0], reverse=True)
        ]

        return {
            "qtdVencidos": qtd_vencidos,
            "qtdVence30": qtd_vence30,
            "qtdVence60": qtd_vence60,
            "qtdSemValidade": qtd_sem_validade,
            "itensCriticos": itens_criticos,
            "porAno": validade_por_ano,
        }

    @staticmethod
    def _build_armas_por_local(depto, seccional, delegacia, status_filtro, modelo, tipo):
        armas_qs = Cautela.objects.filter(
            status=Cautela.STATUS_ATIVA, categoria__iexact="armas"
        ).select_related("item", "item__arma_detalhe")

        if depto:
            armas_qs = armas_qs.filter(depto__iexact=depto)
        if delegacia:
            armas_qs = armas_qs.filter(lotacao__iexact=delegacia)
        if status_filtro:
            armas_qs = armas_qs.filter(item__status__iexact=status_filtro)
        if tipo:
            armas_qs = armas_qs.filter(item__arma_detalhe__tipo__iexact=tipo)
        if modelo:
            armas_qs = armas_qs.filter(item__arma_detalhe__modelo__iexact=modelo)

        armas_por_local = [
            {
                "id": cautela.id,
                "tipo": getattr(getattr(cautela.item, 'arma_detalhe', None), 'tipo', DEFAULT_PLACEHOLDER),
                "modelo": getattr(getattr(cautela.item, 'arma_detalhe', None), 'modelo', DEFAULT_PLACEHOLDER),
                "serie": cautela.item.serie or DEFAULT_PLACEHOLDER, "status": cautela.item.status or DEFAULT_PLACEHOLDER,
                "departamento": cautela.depto or DEFAULT_PLACEHOLDER, "lotacao": cautela.lotacao or DEFAULT_PLACEHOLDER}
            for cautela in armas_qs
        ]

        armas_por_local.sort(key=lambda item: (item["departamento"], item["lotacao"], item["modelo"]))
        return armas_por_local

    def get(self, request):
        categoria = (request.query_params.get("categoria") or "").strip()
        depto = (request.query_params.get("depto") or "").strip()
        seccional = (request.query_params.get("seccional") or "").strip()
        delegacia = (request.query_params.get("delegacia") or "").strip()
        status_filtro = (request.query_params.get("status") or "").strip()
        modelo = (request.query_params.get("modelo") or "").strip()
        tipo = (request.query_params.get("tipo") or "").strip()

        cautelas_filtradas = Cautela.objects.filter(status=Cautela.STATUS_ATIVA)
        if categoria:
            cautelas_filtradas = cautelas_filtradas.filter(item__categoria__icontains=categoria)
        if depto:
            cautelas_filtradas = cautelas_filtradas.filter(depto__iexact=depto)
        if delegacia:
            cautelas_filtradas = cautelas_filtradas.filter(lotacao__iexact=delegacia)
        if status_filtro:
            cautelas_filtradas = cautelas_filtradas.filter(item__status__iexact=status_filtro)
        if tipo:
            cautelas_filtradas = cautelas_filtradas.filter(item__arma_detalhe__tipo__iexact=tipo)
        if modelo:
            cautelas_filtradas = cautelas_filtradas.filter(item__arma_detalhe__modelo__iexact=modelo)

        itens_base = Item.objects.all()
        if categoria:
            itens_base = itens_base.filter(categoria__icontains=categoria)
        if status_filtro:
            itens_base = itens_base.filter(status__iexact=status_filtro)
        if tipo:
            itens_base = itens_base.filter(arma_detalhe__tipo__iexact=tipo)
        if modelo:
            itens_base = itens_base.filter(arma_detalhe__modelo__iexact=modelo)

        # Se um filtro de local for aplicado, filtramos os itens que estão nesses locais via cautela
        if depto or delegacia:
            itens_cautelados_no_local = cautelas_filtradas.values_list('item_id', flat=True)
            itens_base = itens_base.filter(id__in=itens_cautelados_no_local)

        totais = self._build_totais(itens_base)
        categorias = self._build_categorias()
        lotacoes = self._build_locais()
        by_categoria = self._build_ranked_totals(itens_base, "categoria")
        by_status = self._build_ranked_totals(itens_base, "status")
        by_modelo = self._build_by_modelo(itens_base)
        validade = self._build_validade(itens_base)
        tipos = list(
            itens_base
            .filter(categoria__icontains="arma")
            .exclude(arma_detalhe__tipo__isnull=True)
            .exclude(arma_detalhe__tipo__exact="")
            .values_list("arma_detalhe__tipo", flat=True)
            .distinct()
            .order_by("arma_detalhe__tipo")
        )
        modelos = list(
            itens_base
            .filter(categoria__icontains="arma")
            .exclude(arma_detalhe__modelo__isnull=True)
            .exclude(arma_detalhe__modelo__exact="")
            .values_list("arma_detalhe__modelo", flat=True)
            .distinct()
            .order_by("arma_detalhe__modelo")
        )

        ativas = cautelas_filtradas.count()
        registros_cautelas = Cautela.objects.count()
        armas_por_local = self._build_armas_por_local(depto, seccional, delegacia, status_filtro, modelo, tipo)

        return Response(
            {
                "categorias": categorias,
                "locais": lotacoes,
                "tipos": tipos,
                "modelos": modelos,
                "stats": {**totais, "ativas": ativas, "registros_cautelas": registros_cautelas},
                "by_categoria": by_categoria,
                "by_status": by_status,
                "by_modelo": by_modelo,
                "validade": validade,
                "armas_por_local": armas_por_local,
            },
            status=status.HTTP_200_OK,
        )


class BackfillRelationalRefsView(APIView):
    authentication_classes = [SessionAuthentication, BasicAuthentication]
    permission_classes = [IsAdminUser]

    def post(self, request):
        apply_changes = bool(request.data.get("apply", False))
        output = StringIO()

        if apply_changes:
            call_command("backfill_relational_refs", "--apply", stdout=output)
        else:
            call_command("backfill_relational_refs", stdout=output)

        return Response(
            {
                "status": "ok",
                "apply": apply_changes,
                "output": output.getvalue(),
            },
            status=status.HTTP_200_OK,
        )
